"""
backend/report_generator.py
Phase 4: server-side DOCX evidence-pack generation (python-docx).

The client POSTs the session's FULL /api/analyze response plus a REVERSED entity
map (placeholder -> real value). We de-redact every text field and build a
formatted .docx, then stream it back. Nothing is persisted (guardrail #3).

IMPORTANT — data shape: the client sends the full analyze response, where the
substantive analysis is nested under "analysis" and judgment/attestation are
siblings:
    { "analysis": { executive_summary, overall_severity, red_flags[],
                    recommended_actions[], mom_report_draft{subject,to,body} },
      "judgment": {...}, "attestation": {...}, "entity_map": {...} }
"""
import io


def _deRedact(text: str, emap_reversed: dict) -> str:
    """Replace every [PLACEHOLDER_N] token with its real value.

    Longest placeholders first so a longer token isn't partially clobbered by a
    shorter one. emap_reversed is {placeholder -> real}.
    """
    if not text or not emap_reversed:
        return text or ""
    for placeholder, real in sorted(emap_reversed.items(), key=lambda x: len(x[0]), reverse=True):
        text = text.replace(placeholder, real)
    return text


def generate_docx(report: dict, emap_reversed: dict, filenames: list, generated_at: str) -> bytes:
    """Build the DOCX evidence pack. `report` is the full /api/analyze response."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # The substantive analysis is nested; judgment + attestation are siblings.
    inner = report.get("analysis") or {}
    judgment = report.get("judgment") or {}
    attestation = report.get("attestation") or {}

    def dr(t):
        return _deRedact(t if isinstance(t, str) else ("" if t is None else str(t)), emap_reversed)

    doc = Document()

    # ── HEADER ──────────────────────────────────────────────────────────────
    title = doc.add_heading("ClauseGuard Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {generated_at}\n").italic = True
    meta.add_run(f"Documents analysed: {', '.join(filenames) if filenames else '—'}\n").italic = True
    report_hash = attestation.get("report_hash")
    if report_hash:
        meta.add_run(f"Report hash (SHA-256): {report_hash}\n").italic = True
    meta.add_run(
        "This report is not legal advice and is not exhaustive. "
        "A 'no red flags found' result does not guarantee a contract is fair. "
        "Singapore employment contracts (MVP scope)."
    ).italic = True
    doc.add_paragraph()

    # ── DISPUTE JUDGMENT (only when a real verdict exists) ──────────────────
    verdict = (judgment.get("verdict") or "").strip()
    if verdict and verdict != "INSUFFICIENT_INFORMATION":
        doc.add_heading("Dispute Judgment", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Verdict: {verdict}  |  Confidence: {judgment.get('confidence', '')}").bold = True
        ds = dr(judgment.get("dispute_summary", ""))
        if ds:
            doc.add_paragraph(ds)
        vr = dr(judgment.get("verdict_reasoning", ""))
        if vr:
            doc.add_paragraph(vr)
        forum = judgment.get("recommended_forum")
        if forum:
            fp = doc.add_paragraph()
            fp.add_run("Recommended forum: ").bold = True
            fp.add_run(dr(f"{forum} — {judgment.get('forum_reasoning', '')}".strip(" —")))

    # ── EXECUTIVE SUMMARY ───────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(dr(inner.get("executive_summary") or "No summary available."))
    sev = inner.get("overall_severity")
    if sev:
        sp = doc.add_paragraph()
        sp.add_run("Overall severity: ").bold = True
        sp.add_run(str(sev))

    # ── RED FLAGS (capped at 20 for readability) ────────────────────────────
    doc.add_heading("Red Flags", level=1)
    flags = (inner.get("red_flags") or [])[:20]
    if flags:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Severity", "Issue", "Clause / Document"
        for f in flags:
            row = table.add_row().cells
            row[0].text = str(f.get("severity", ""))
            title_line = (f.get("title") or "").strip()
            issue = (f.get("issue") or "").strip()
            if title_line and issue and title_line != issue:
                body_txt = f"{title_line}: {issue}"
            else:
                body_txt = issue or title_line
            row[1].text = dr(body_txt)
            row[2].text = dr(f.get("clause_or_section") or f.get("document") or "")
    else:
        doc.add_paragraph("No red flags identified.")
    doc.add_paragraph()

    # ── RECOMMENDED ACTIONS (list of dicts) ─────────────────────────────────
    doc.add_heading("Recommended Actions", level=1)
    actions = inner.get("recommended_actions") or []
    if isinstance(actions, list) and actions:
        for a in actions:
            if isinstance(a, dict):
                txt = dr(a.get("action", ""))
                meta_bits = [b for b in (a.get("channel"), a.get("urgency")) if b]
                if meta_bits:
                    txt += f" ({' · '.join(str(b) for b in meta_bits)})"
                if a.get("notes"):
                    txt += f" — {dr(a['notes'])}"
            else:
                txt = dr(str(a))
            doc.add_paragraph(txt, style="List Bullet")
    else:
        doc.add_paragraph("None provided.")

    # ── MOM / TADM DRAFT LETTER (mom_report_draft is a dict) ────────────────
    doc.add_heading("MOM / TADM Draft Letter", level=1)
    doc.add_paragraph(
        "Fill in any remaining [BRACKETED] fields before sending. "
        "This letter was generated from an AI analysis — review carefully before submission."
    ).italic = True
    draft = inner.get("mom_report_draft")
    if isinstance(draft, dict):
        doc.add_paragraph(f"To: {dr(draft.get('to') or 'MOM Contact Centre / TADM')}")
        doc.add_paragraph(f"Subject: {dr(draft.get('subject', ''))}")
        doc.add_paragraph()
        for line in dr(draft.get("body", "")).split("\n"):
            doc.add_paragraph(line)
    elif draft:
        for line in dr(str(draft)).split("\n"):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("Not generated.")

    # ── ATTESTATION RECEIPT ─────────────────────────────────────────────────
    doc.add_heading("Attestation Receipt", level=1)
    doc.add_paragraph(
        "The hash below is a SHA-256 fingerprint of the original analysis report. The "
        "signature is an HMAC (symmetric) signature via Terminal 3: it proves the analysis "
        "was not altered after signing, but it does NOT prove the analysis is correct, and "
        "third parties cannot independently verify it without the signing key."
    ).italic = True
    if attestation:
        for key, val in attestation.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(val))
    else:
        doc.add_paragraph("No attestation data available for this session.")

    # ── FOOTER ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph(
        "ClauseGuard · Not legal advice · probono.sg for free legal consultation · "
        "MOM: mom.gov.sg · TADM: tadm.gov.sg"
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
